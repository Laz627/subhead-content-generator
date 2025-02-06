import streamlit as st
import openai
import requests
import re
from urllib.parse import quote_plus
from docx import Document
from io import BytesIO
import time

# ----------------------------
# Page Configuration & Title
# ----------------------------
st.set_page_config(page_title="SEO Element Generator", layout="wide")
st.title("SEO Element Generator")
st.write("Created by Brandon Lazovic")

st.markdown("""
## How to use this app:
1. Enter your OpenAI API key.
2. Enter your DataForSEO credentials.
3. Input up to 10 target keywords (one per line).
4. Click **Generate SEO Elements** to get recommendations.
5. Review the results and explanations.
6. Download the results as a Word document.
""")

# ----------------------------
# User Inputs for Credentials and Keywords
# ----------------------------
openai_api_key = st.text_input("Enter your OpenAI API key:", type="password")
dataforseo_username = st.text_input("Enter your DataForSEO username:")
dataforseo_password = st.text_input("Enter your DataForSEO password:", type="password")
keywords = st.text_area("Enter up to 10 target keywords (one per line):", height=200)
keyword_list = [k.strip() for k in keywords.split("\n") if k.strip()]

# ----------------------------
# Function: DataForSEO Google SERP Scraper
# ----------------------------
def scrape_google_results(keyword, username, password, num_results=10):
    url = "https://api.dataforseo.com/v3/serp/google/organic/live/advanced"
    payload = [{
        "keyword": keyword,
        "language_code": "en",   # Adjust if needed
        "location_code": 2840,   # Example: 2840 corresponds to the United States
        "device": "desktop",     # Options: "desktop" or "mobile"
        "num": num_results
    }]
    
    response = requests.post(url, auth=(username, password), json=payload)
    data = response.json()
    
    results = []
    for task in data.get("tasks", []):
        for result_item in task.get("result", []):
            for item in result_item.get("items", []):
                title = item.get("title", "")
                snippet = item.get("snippet", "")
                if title:
                    results.append({"title": title, "snippet": snippet})
    return results[:num_results]

# ----------------------------
# Function: Summarize Competitor Elements
# ----------------------------
def summarize_competitor_elements(results):
    if not results:
        return "No competitor results found. Unable to perform competitor analysis."
    
    titles = [result["title"] for result in results]
    snippets = [result["snippet"] for result in results]
    
    avg_title_length = sum(len(title) for title in titles) / len(titles)
    avg_snippet_length = sum(len(snippet) for snippet in snippets) / len(snippets)
    
    summary = f"Analyzed {len(results)} competitor results.\n"
    summary += f"Average title length: {avg_title_length:.1f} characters.\n"
    summary += f"Average snippet length: {avg_snippet_length:.1f} characters.\n"
    
    # Count word frequencies in titles (ignoring words shorter than 4 characters)
    word_freq = {}
    for title in titles:
        for word in re.findall(r'\w+', title.lower()):
            if len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1
    common_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
    summary += f"Common words in titles: {', '.join([word for word, _ in common_words])}\n\n"
    
    summary += "Sample competitor titles:\n"
    for title in titles[:5]:
        summary += f"- {title}\n"
    
    summary += "\nSample competitor snippets:\n"
    for snippet in snippets[:5]:
        summary += f"- {snippet[:100]}...\n"
    
    return summary

# ----------------------------
# Function: Generate SEO Elements using OpenAI API
# ----------------------------
def generate_seo_elements(keyword, competitor_summary, openai_api_key, max_retries=3):
    # Set the OpenAI API key for the session
    openai.api_key = openai_api_key

    prompt = f"""
Generate an H1, title tag, and meta description for the keyword: "{keyword}"

Requirements:
- H1 and title tag should be 70 characters or less.
- Meta description should be 155 characters or less.
- Avoid buzzwords and branded terms.
- Include an exact match or close variation of the target keyword.
- Closely align with the competitor results provided below.
- The elements should be a summarization of common elements from the top 10 competitors.

Competitor analysis:
{competitor_summary}

Based on the competitor analysis, create SEO elements that are very similar to the competitors' approach, while still being unique. Focus on common phrases, structures, and themes used by competitors.

Please provide your response in the following structure:
COMPETITOR ELEMENTS SUMMARY
1. Most Common Title Structures:
   - [Point 1]
   - [Point 2]
   ...
2. Common Themes in Meta Descriptions:
   - [Point 1]
   - [Point 2]
   ...

3. Frequently Used Phrases or Keywords:
   - [Phrase 1]
   - [Phrase 2]
   ...
4. Notable Patterns in Competitor Information Presentation:
   - [Pattern 1]
   - [Pattern 2]
   ...

SEO ELEMENTS
H1: [Your H1]
Explanation:
- [Point 1]
- [Point 2]
...
Title Tag: [Your Title Tag]
Explanation:
- [Point 1]
- [Point 2]
...
Meta Description: [Your Meta Description]
Explanation:
- [Point 1]
- [Point 2]
...
    """
    
    for attempt in range(max_retries):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an SEO expert tasked with creating optimized on-page elements that closely align with competitor trends."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            if attempt == max_retries - 1:
                raise
            time.sleep(2 ** attempt)  # Exponential backoff

# ----------------------------
# Function: Create Word Document
# ----------------------------
def create_word_document(results):
    doc = Document()
    doc.add_heading('SEO Element Generator Results', 0)
    
    for result in results:
        doc.add_heading(f"Keyword: {result['Keyword']}", level=1)
        doc.add_paragraph(result['SEO Elements and Competitor Summary'])
        doc.add_heading("Competitor Analysis", level=2)
        doc.add_paragraph(result['Competitor Analysis'])
        doc.add_paragraph("\n")  # Blank line between results
    
    return doc

# ----------------------------
# Main Application Logic
# ----------------------------
if st.button("Generate SEO Elements") and openai_api_key and dataforseo_username and dataforseo_password and keyword_list:
    results = []
    
    for keyword in keyword_list[:10]:
        st.subheader(f"Results for: {keyword}")
        
        with st.spinner(f"Analyzing competitors for '{keyword}'..."):
            competitor_results = scrape_google_results(keyword, dataforseo_username, dataforseo_password)
            competitor_summary = summarize_competitor_elements(competitor_results)
        
        with st.spinner(f"Generating SEO elements for '{keyword}'..."):
            seo_elements = generate_seo_elements(keyword, competitor_summary, openai_api_key)
        
        st.write(seo_elements)
        st.write("Competitor Analysis Summary:")
        st.write(competitor_summary)
        
        results.append({
            "Keyword": keyword,
            "SEO Elements and Competitor Summary": seo_elements,
            "Competitor Analysis": competitor_summary
        })
    
    # Create and offer the Word document download
    doc = create_word_document(results)
    bio = BytesIO()
    doc.save(bio)
    st.download_button(
        label="Download results as Word Document",
        data=bio.getvalue(),
        file_name="seo_elements_results.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.write("Please enter your OpenAI API key, DataForSEO credentials, and at least one keyword to generate SEO elements.")
